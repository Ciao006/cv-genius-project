"""
Multi-Format CV Export Service
Free and open-source CV export to PDF, DOCX, TXT, and HTML formats
"""

import asyncio
import io
import json
import logging
from typing import Dict, List, Any, Optional, Union, BinaryIO
from dataclasses import dataclass
from enum import Enum
from datetime import datetime
import tempfile
import os

# For PDF generation
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False

# For DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# For HTML templating
try:
    from jinja2 import Environment, FileSystemLoader, Template
    HAS_JINJA2 = True
except ImportError:
    HAS_JINJA2 = False

logger = logging.getLogger(__name__)

class ExportFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    JSON = "json"

class ExportQuality(Enum):
    DRAFT = "draft"
    STANDARD = "standard"
    HIGH = "high"
    PRINT = "print"

@dataclass
class ExportSettings:
    format: ExportFormat
    quality: ExportQuality = ExportQuality.STANDARD
    template: str = "modern_tech"
    include_contact_info: bool = True
    include_photo: bool = False
    page_size: str = "A4"  # A4, Letter, Legal
    margins: Dict[str, float] = None  # top, bottom, left, right in mm
    font_family: str = "Arial"
    font_size: int = 11
    line_spacing: float = 1.15
    compress: bool = True
    watermark: Optional[str] = None
    
    def __post_init__(self):
        if self.margins is None:
            self.margins = {"top": 20, "bottom": 20, "left": 20, "right": 20}

@dataclass
class ExportResult:
    success: bool
    file_data: Optional[bytes] = None
    file_path: Optional[str] = None
    filename: str = ""
    content_type: str = ""
    file_size: int = 0
    error_message: Optional[str] = None
    generation_time_ms: int = 0
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}

class PDFExporter:
    """
    PDF export using WeasyPrint
    """
    
    def __init__(self):
        self.font_config = FontConfiguration() if HAS_WEASYPRINT else None
    
    def export(self, cv_data: Dict[str, Any], template_html: str, 
              settings: ExportSettings) -> ExportResult:
        """Export CV to PDF format"""
        
        if not HAS_WEASYPRINT:
            return ExportResult(
                success=False,
                error_message="WeasyPrint not installed. Install with: pip install weasyprint"
            )
        
        try:
            start_time = datetime.now()
            
            # Generate CSS based on settings
            css_content = self._generate_pdf_css(settings)
            
            # Create HTML object
            html_doc = HTML(string=template_html)
            css_doc = CSS(string=css_content, font_config=self.font_config)
            
            # Generate PDF
            pdf_buffer = io.BytesIO()
            html_doc.write_pdf(pdf_buffer, stylesheets=[css_doc], font_config=self.font_config)
            
            pdf_data = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            generation_time = int((datetime.now() - start_time).total_seconds() * 1000)
            
            # Generate filename
            name = cv_data.get("personal_details", {}).get("full_name", "CV")
            filename = f"{name.replace(' ', '_')}_CV.pdf"
            
            return ExportResult(
                success=True,
                file_data=pdf_data,
                filename=filename,
                content_type="application/pdf",
                file_size=len(pdf_data),
                generation_time_ms=generation_time,
                metadata={
                    "page_count": self._get_pdf_page_count(pdf_data),
                    "template": settings.template,
                    "quality": settings.quality.value
                }
            )
            
        except Exception as e:
            logger.error(f"PDF export error: {e}")
            return ExportResult(
                success=False,
                error_message=f"PDF generation failed: {str(e)}"
            )
    
    def _generate_pdf_css(self, settings: ExportSettings) -> str:
        """Generate CSS for PDF styling"""
        
        page_size = settings.page_size.lower()
        margins = settings.margins
        
        css = f"""
        @page {{
            size: {page_size};
            margin-top: {margins['top']}mm;
            margin-bottom: {margins['bottom']}mm;
            margin-left: {margins['left']}mm;
            margin-right: {margins['right']}mm;
        }}
        
        body {{
            font-family: '{settings.font_family}', Arial, sans-serif;
            font-size: {settings.font_size}pt;
            line-height: {settings.line_spacing};
            color: #000;
            margin: 0;
            padding: 0;
        }}
        
        .template-container {{
            width: 100%;
            max-width: none;
            margin: 0;
            box-shadow: none;
        }}
        
        /* Print-specific styles */
        * {{
            print-color-adjust: exact;
            -webkit-print-color-adjust: exact;
        }}
        
        .no-break {{
            page-break-inside: avoid;
        }}
        
        .page-break {{
            page-break-before: always;
        }}
        
        h1, h2, h3 {{
            page-break-after: avoid;
        }}
        
        /* Ensure colors print correctly */
        .header {{
            background: #3182ce !important;
            color: white !important;
        }}
        
        .skill-item {{
            background: #3182ce !important;
            color: white !important;
        }}
        """
        
        # Quality-specific adjustments
        if settings.quality == ExportQuality.HIGH or settings.quality == ExportQuality.PRINT:
            css += """
            body {
                font-size: 12pt;
                line-height: 1.3;
            }
            
            .section-title {
                font-size: 14pt;
            }
            
            .header-name {
                font-size: 24pt;
            }
            """
        
        # Add watermark if specified
        if settings.watermark:
            css += f"""
            @page {{
                background-image: url("data:image/svg+xml;charset=utf-8,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 400 200'%3E%3Ctext x='50%25' y='50%25' text-anchor='middle' font-size='40' fill='%23f0f0f0' opacity='0.3' transform='rotate(-45 200 100)'%3E{settings.watermark}%3C/text%3E%3C/svg%3E");
                background-repeat: repeat;
                background-position: center;
            }}
            """
        
        return css
    
    def _get_pdf_page_count(self, pdf_data: bytes) -> int:
        """Get number of pages in PDF (simplified implementation)"""
        try:
            # This is a very basic implementation
            # In practice, you'd use a proper PDF library like PyPDF2
            page_count = pdf_data.count(b'/Type /Page')
            return max(1, page_count)
        except:
            return 1

class DOCXExporter:
    """
    DOCX export using python-docx
    """
    
    def __init__(self):
        pass
    
    def export(self, cv_data: Dict[str, Any], template_html: str, 
              settings: ExportSettings) -> ExportResult:
        """Export CV to DOCX format"""
        
        if not HAS_DOCX:
            return ExportResult(
                success=False,
                error_message="python-docx not installed. Install with: pip install python-docx"
            )
        
        try:
            start_time = datetime.now()
            
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_docx_styles(doc, settings)
            
            # Add content
            self._add_header_section(doc, cv_data.get("personal_details", {}), settings)
            
            if cv_data.get("professional_summary"):
                self._add_summary_section(doc, cv_data["professional_summary"])
            
            if cv_data.get("skills"):
                self._add_skills_section(doc, cv_data["skills"])
            
            if cv_data.get("work_experience"):
                self._add_experience_section(doc, cv_data["work_experience"])
            
            if cv_data.get("education"):
                self._add_education_section(doc, cv_data["education"])
            
            if cv_data.get("projects"):
                self._add_projects_section(doc, cv_data["projects"])
            
            # Save to buffer
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            generation_time = int((datetime.now() - start_time).total_seconds() * 1000)
            
            # Generate filename
            name = cv_data.get("personal_details", {}).get("full_name", "CV")
            filename = f"{name.replace(' ', '_')}_CV.docx"
            
            return ExportResult(
                success=True,
                file_data=docx_data,
                filename=filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_size=len(docx_data),
                generation_time_ms=generation_time,
                metadata={
                    "page_count": 1,  # Approximate
                    "template": settings.template,
                    "word_count": self._estimate_word_count(cv_data)
                }
            )
            
        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            return ExportResult(
                success=False,
                error_message=f"DOCX generation failed: {str(e)}"
            )
    
    def _setup_docx_styles(self, doc: Document, settings: ExportSettings):
        """Set up document styles"""
        
        # Document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(settings.margins["top"] / 25.4)
            section.bottom_margin = Inches(settings.margins["bottom"] / 25.4)
            section.left_margin = Inches(settings.margins["left"] / 25.4)
            section.right_margin = Inches(settings.margins["right"] / 25.4)
        
        # Create custom styles
        styles = doc.styles
        
        # Header style
        if 'CV Header' not in [style.name for style in styles]:
            header_style = styles.add_style('CV Header', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = settings.font_family
            header_font.size = Pt(18)
            header_font.bold = True
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Section heading style
        if 'CV Section' not in [style.name for style in styles]:
            section_style = styles.add_style('CV Section', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = settings.font_family
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = None  # Blue color would be set here
        
        # Body style
        if 'CV Body' not in [style.name for style in styles]:
            body_style = styles.add_style('CV Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = settings.font_family
            body_font.size = Pt(settings.font_size)
    
    def _add_header_section(self, doc: Document, personal_details: Dict[str, Any], 
                           settings: ExportSettings):
        """Add header section with personal details"""
        
        # Name
        name_para = doc.add_paragraph(personal_details.get("full_name", ""), style='CV Header')
        
        # Title
        if personal_details.get("desired_position"):
            title_para = doc.add_paragraph(personal_details["desired_position"])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].font.size = Pt(14)
            title_para.runs[0].font.italic = True
        
        # Contact info
        if settings.include_contact_info:
            contact_items = []
            if personal_details.get("phone"):
                contact_items.append(personal_details["phone"])
            if personal_details.get("email"):
                contact_items.append(personal_details["email"])
            if personal_details.get("location"):
                contact_items.append(personal_details["location"])
            if personal_details.get("linkedin_url"):
                contact_items.append(personal_details["linkedin_url"])
            
            if contact_items:
                contact_para = doc.add_paragraph(" | ".join(contact_items))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para.runs[0].font.size = Pt(10)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_summary_section(self, doc: Document, summary: str):
        """Add professional summary section"""
        
        doc.add_paragraph("PROFESSIONAL SUMMARY", style='CV Section')
        doc.add_paragraph(summary, style='CV Body')
        doc.add_paragraph()
    
    def _add_skills_section(self, doc: Document, skills: Dict[str, List[str]]):
        """Add skills section"""
        
        doc.add_paragraph("TECHNICAL SKILLS", style='CV Section')
        
        for category, skill_list in skills.items():
            if skill_list:
                skills_para = doc.add_paragraph(style='CV Body')
                skills_para.add_run(f"{category.title()}: ").font.bold = True
                skills_para.add_run(", ".join(skill_list))
        
        doc.add_paragraph()
    
    def _add_experience_section(self, doc: Document, experience: List[Dict[str, Any]]):
        """Add work experience section"""
        
        doc.add_paragraph("WORK EXPERIENCE", style='CV Section')
        
        for exp in experience:
            # Job header
            job_para = doc.add_paragraph(style='CV Body')
            job_para.add_run(exp.get("job_title", "")).font.bold = True
            job_para.add_run(f" | {exp.get('company', '')}")
            
            # Dates
            date_para = doc.add_paragraph(style='CV Body')
            start_date = exp.get("start_date", "")
            end_date = exp.get("end_date", "Present") if not exp.get("is_current") else "Present"
            date_para.add_run(f"{start_date} - {end_date}").font.italic = True
            
            # Location
            if exp.get("location"):
                date_para.add_run(f" | {exp['location']}").font.italic = True
            
            # Achievements
            if exp.get("achievements"):
                for achievement in exp["achievements"]:
                    achievement_para = doc.add_paragraph(f"• {achievement}", style='CV Body')
                    achievement_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()
    
    def _add_education_section(self, doc: Document, education: List[Dict[str, Any]]):
        """Add education section"""
        
        doc.add_paragraph("EDUCATION", style='CV Section')
        
        for edu in education:
            edu_para = doc.add_paragraph(style='CV Body')
            edu_para.add_run(edu.get("degree", "")).font.bold = True
            edu_para.add_run(f" | {edu.get('institution', '')}")
            
            # Dates and grade
            details = []
            if edu.get("start_date") or edu.get("end_date"):
                start = edu.get("start_date", "")
                end = edu.get("end_date", "")
                if start or end:
                    details.append(f"{start} - {end}" if start and end else start or end)
            
            if edu.get("grade"):
                details.append(edu["grade"])
            
            if details:
                details_para = doc.add_paragraph(" | ".join(details), style='CV Body')
                details_para.runs[0].font.italic = True
        
        doc.add_paragraph()
    
    def _add_projects_section(self, doc: Document, projects: List[Dict[str, Any]]):
        """Add projects section"""
        
        doc.add_paragraph("KEY PROJECTS", style='CV Section')
        
        for project in projects:
            project_para = doc.add_paragraph(style='CV Body')
            project_para.add_run(project.get("name", "")).font.bold = True
            
            if project.get("description"):
                desc_para = doc.add_paragraph(project["description"], style='CV Body')
                desc_para.paragraph_format.left_indent = Inches(0.25)
            
            if project.get("technologies"):
                tech_para = doc.add_paragraph(style='CV Body')
                tech_para.paragraph_format.left_indent = Inches(0.25)
                tech_para.add_run("Technologies: ").font.italic = True
                tech_para.add_run(", ".join(project["technologies"]))
        
        doc.add_paragraph()
    
    def _estimate_word_count(self, cv_data: Dict[str, Any]) -> int:
        """Estimate word count in CV"""
        word_count = 0
        
        # Count words in text fields
        text_fields = [
            cv_data.get("professional_summary", "")
        ]
        
        if cv_data.get("work_experience"):
            for exp in cv_data["work_experience"]:
                text_fields.extend(exp.get("achievements", []))
        
        for text in text_fields:
            word_count += len(str(text).split())
        
        return word_count

class TXTExporter:
    """
    Plain text export
    """
    
    def export(self, cv_data: Dict[str, Any], template_html: str, 
              settings: ExportSettings) -> ExportResult:
        """Export CV to plain text format"""
        
        try:
            start_time = datetime.now()
            
            # Build text content
            text_lines = []
            
            # Header
            personal = cv_data.get("personal_details", {})
            if personal.get("full_name"):
                text_lines.append(personal["full_name"].upper())
                text_lines.append("=" * len(personal["full_name"]))
            
            if personal.get("desired_position"):
                text_lines.append(personal["desired_position"])
            
            # Contact info
            if settings.include_contact_info:
                contact_info = []
                if personal.get("phone"):
                    contact_info.append(f"Phone: {personal['phone']}")
                if personal.get("email"):
                    contact_info.append(f"Email: {personal['email']}")
                if personal.get("location"):
                    contact_info.append(f"Location: {personal['location']}")
                if personal.get("linkedin_url"):
                    contact_info.append(f"LinkedIn: {personal['linkedin_url']}")
                
                if contact_info:
                    text_lines.extend(contact_info)
            
            text_lines.append("")
            
            # Professional Summary
            if cv_data.get("professional_summary"):
                text_lines.append("PROFESSIONAL SUMMARY")
                text_lines.append("-" * 20)
                text_lines.append(cv_data["professional_summary"])
                text_lines.append("")
            
            # Skills
            if cv_data.get("skills"):
                text_lines.append("TECHNICAL SKILLS")
                text_lines.append("-" * 16)
                for category, skill_list in cv_data["skills"].items():
                    if skill_list:
                        text_lines.append(f"{category.title()}: {', '.join(skill_list)}")
                text_lines.append("")
            
            # Work Experience
            if cv_data.get("work_experience"):
                text_lines.append("WORK EXPERIENCE")
                text_lines.append("-" * 15)
                
                for exp in cv_data["work_experience"]:
                    # Job header
                    job_title = exp.get("job_title", "")
                    company = exp.get("company", "")
                    text_lines.append(f"{job_title} | {company}")
                    
                    # Dates and location
                    date_info = []
                    start_date = exp.get("start_date", "")
                    end_date = exp.get("end_date", "Present") if not exp.get("is_current") else "Present"
                    if start_date:
                        date_info.append(f"{start_date} - {end_date}")
                    if exp.get("location"):
                        date_info.append(exp["location"])
                    
                    if date_info:
                        text_lines.append(" | ".join(date_info))
                    
                    # Achievements
                    if exp.get("achievements"):
                        for achievement in exp["achievements"]:
                            text_lines.append(f"• {achievement}")
                    
                    text_lines.append("")
            
            # Education
            if cv_data.get("education"):
                text_lines.append("EDUCATION")
                text_lines.append("-" * 9)
                
                for edu in cv_data["education"]:
                    degree = edu.get("degree", "")
                    institution = edu.get("institution", "")
                    text_lines.append(f"{degree} | {institution}")
                    
                    details = []
                    if edu.get("start_date") or edu.get("end_date"):
                        start = edu.get("start_date", "")
                        end = edu.get("end_date", "")
                        if start or end:
                            details.append(f"{start} - {end}" if start and end else start or end)
                    
                    if edu.get("grade"):
                        details.append(edu["grade"])
                    
                    if details:
                        text_lines.append(" | ".join(details))
                    
                    text_lines.append("")
            
            # Projects
            if cv_data.get("projects"):
                text_lines.append("KEY PROJECTS")
                text_lines.append("-" * 12)
                
                for project in cv_data["projects"]:
                    text_lines.append(project.get("name", ""))
                    
                    if project.get("description"):
                        text_lines.append(f"Description: {project['description']}")
                    
                    if project.get("technologies"):
                        text_lines.append(f"Technologies: {', '.join(project['technologies'])}")
                    
                    text_lines.append("")
            
            # Join all lines
            text_content = "\n".join(text_lines)
            text_data = text_content.encode('utf-8')
            
            generation_time = int((datetime.now() - start_time).total_seconds() * 1000)
            
            # Generate filename
            name = cv_data.get("personal_details", {}).get("full_name", "CV")
            filename = f"{name.replace(' ', '_')}_CV.txt"
            
            return ExportResult(
                success=True,
                file_data=text_data,
                filename=filename,
                content_type="text/plain",
                file_size=len(text_data),
                generation_time_ms=generation_time,
                metadata={
                    "line_count": len(text_lines),
                    "character_count": len(text_content),
                    "word_count": len(text_content.split())
                }
            )
            
        except Exception as e:
            logger.error(f"TXT export error: {e}")
            return ExportResult(
                success=False,
                error_message=f"TXT generation failed: {str(e)}"
            )

class HTMLExporter:
    """
    HTML export with styling
    """
    
    def export(self, cv_data: Dict[str, Any], template_html: str, 
              settings: ExportSettings) -> ExportResult:
        """Export CV to HTML format"""
        
        try:
            start_time = datetime.now()
            
            # Add inline CSS for standalone HTML
            standalone_html = self._make_standalone_html(template_html, settings)
            
            html_data = standalone_html.encode('utf-8')
            
            generation_time = int((datetime.now() - start_time).total_seconds() * 1000)
            
            # Generate filename
            name = cv_data.get("personal_details", {}).get("full_name", "CV")
            filename = f"{name.replace(' ', '_')}_CV.html"
            
            return ExportResult(
                success=True,
                file_data=html_data,
                filename=filename,
                content_type="text/html",
                file_size=len(html_data),
                generation_time_ms=generation_time,
                metadata={
                    "template": settings.template,
                    "includes_css": True,
                    "standalone": True
                }
            )
            
        except Exception as e:
            logger.error(f"HTML export error: {e}")
            return ExportResult(
                success=False,
                error_message=f"HTML generation failed: {str(e)}"
            )
    
    def _make_standalone_html(self, template_html: str, settings: ExportSettings) -> str:
        """Make HTML standalone by inlining CSS"""
        
        # Add meta tags and ensure proper structure
        if not template_html.startswith('<!DOCTYPE'):
            template_html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CV Export</title>
</head>
<body>
{template_html}
</body>
</html>'''
        
        # Add export-specific styles
        export_css = """
        <style>
        /* Export optimizations */
        * {
            -webkit-print-color-adjust: exact;
            print-color-adjust: exact;
        }
        
        @media print {
            body {
                margin: 0;
                padding: 0;
            }
            
            .template-container {
                box-shadow: none;
                margin: 0;
            }
        }
        
        .export-note {
            display: none;
        }
        </style>
        """
        
        # Insert CSS before closing head tag
        if '</head>' in template_html:
            template_html = template_html.replace('</head>', f'{export_css}</head>')
        
        return template_html

class JSONExporter:
    """
    JSON export for data portability
    """
    
    def export(self, cv_data: Dict[str, Any], template_html: str, 
              settings: ExportSettings) -> ExportResult:
        """Export CV to JSON format"""
        
        try:
            start_time = datetime.now()
            
            # Clean and structure data for export
            export_data = {
                "cv_data": cv_data,
                "export_metadata": {
                    "export_date": datetime.now().isoformat(),
                    "export_format": "json",
                    "export_settings": {
                        "template": settings.template,
                        "quality": settings.quality.value,
                        "include_contact_info": settings.include_contact_info
                    },
                    "version": "1.0"
                }
            }
            
            # Convert to JSON
            json_content = json.dumps(export_data, indent=2, ensure_ascii=False)
            json_data = json_content.encode('utf-8')
            
            generation_time = int((datetime.now() - start_time).total_seconds() * 1000)
            
            # Generate filename
            name = cv_data.get("personal_details", {}).get("full_name", "CV")
            filename = f"{name.replace(' ', '_')}_CV.json"
            
            return ExportResult(
                success=True,
                file_data=json_data,
                filename=filename,
                content_type="application/json",
                file_size=len(json_data),
                generation_time_ms=generation_time,
                metadata={
                    "structure_version": "1.0",
                    "data_sections": list(cv_data.keys()),
                    "json_size_kb": round(len(json_data) / 1024, 2)
                }
            )
            
        except Exception as e:
            logger.error(f"JSON export error: {e}")
            return ExportResult(
                success=False,
                error_message=f"JSON generation failed: {str(e)}"
            )

class ExportService:
    """
    Main export service that coordinates all export formats
    """
    
    def __init__(self):
        self.exporters = {
            ExportFormat.PDF: PDFExporter(),
            ExportFormat.DOCX: DOCXExporter(),
            ExportFormat.TXT: TXTExporter(),
            ExportFormat.HTML: HTMLExporter(),
            ExportFormat.JSON: JSONExporter()
        }
        
        # Template loader
        self.template_loader = None
        if HAS_JINJA2:
            template_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
            if os.path.exists(template_dir):
                self.template_loader = Environment(loader=FileSystemLoader(template_dir))
    
    async def export_cv(self, cv_data: Dict[str, Any], settings: ExportSettings) -> ExportResult:
        """Export CV in the specified format"""
        
        try:
            # Get the appropriate exporter
            exporter = self.exporters.get(settings.format)
            if not exporter:
                return ExportResult(
                    success=False,
                    error_message=f"Unsupported export format: {settings.format.value}"
                )
            
            # Generate template HTML if needed
            template_html = ""
            if settings.format in [ExportFormat.PDF, ExportFormat.HTML]:
                template_html = await self._generate_template_html(cv_data, settings.template)
                if not template_html:
                    return ExportResult(
                        success=False,
                        error_message=f"Failed to load template: {settings.template}"
                    )
            
            # Perform export
            result = exporter.export(cv_data, template_html, settings)
            
            # Add common metadata
            if result.success and result.metadata:
                result.metadata.update({
                    "export_format": settings.format.value,
                    "export_timestamp": datetime.now().isoformat(),
                    "cv_id": cv_data.get("cv_id", "unknown"),
                    "user_id": cv_data.get("user_id", "unknown")
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Export service error: {e}")
            return ExportResult(
                success=False,
                error_message=f"Export failed: {str(e)}"
            )
    
    async def batch_export(self, cv_data: Dict[str, Any], 
                          formats: List[ExportFormat],
                          settings: ExportSettings = None) -> Dict[str, ExportResult]:
        """Export CV in multiple formats"""
        
        if settings is None:
            settings = ExportSettings(format=formats[0])
        
        results = {}
        
        # Export each format
        for format_type in formats:
            format_settings = ExportSettings(
                format=format_type,
                quality=settings.quality,
                template=settings.template,
                include_contact_info=settings.include_contact_info,
                include_photo=settings.include_photo,
                page_size=settings.page_size,
                margins=settings.margins,
                font_family=settings.font_family,
                font_size=settings.font_size,
                line_spacing=settings.line_spacing,
                compress=settings.compress,
                watermark=settings.watermark
            )
            
            result = await self.export_cv(cv_data, format_settings)
            results[format_type.value] = result
        
        return results
    
    async def _generate_template_html(self, cv_data: Dict[str, Any], template_name: str) -> str:
        """Generate HTML from template"""
        
        try:
            if self.template_loader:
                # Use Jinja2 template
                template = self.template_loader.get_template(f"{template_name}.html")
                return template.render(**cv_data)
            else:
                # Fallback: load template file directly
                template_path = os.path.join(
                    os.path.dirname(__file__), '..', 'templates', f"{template_name}.html"
                )
                
                if os.path.exists(template_path):
                    with open(template_path, 'r', encoding='utf-8') as f:
                        template_content = f.read()
                    
                    # Simple template variable replacement
                    return self._simple_template_render(template_content, cv_data)
                else:
                    # Generate basic HTML
                    return self._generate_basic_html(cv_data)
        
        except Exception as e:
            logger.error(f"Template generation error: {e}")
            return self._generate_basic_html(cv_data)
    
    def _simple_template_render(self, template_content: str, cv_data: Dict[str, Any]) -> str:
        """Simple template variable replacement"""
        
        # Replace basic variables
        for key, value in cv_data.items():
            if isinstance(value, dict):
                for subkey, subvalue in value.items():
                    template_content = template_content.replace(
                        f"{{{{ {key}.{subkey} }}}}", str(subvalue or "")
                    )
            else:
                template_content = template_content.replace(
                    f"{{{{ {key} }}}}", str(value or "")
                )
        
        return template_content
    
    def _generate_basic_html(self, cv_data: Dict[str, Any]) -> str:
        """Generate basic HTML structure"""
        
        personal = cv_data.get("personal_details", {})
        
        html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{personal.get('full_name', 'CV')}</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    line-height: 1.6;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                    padding: 20px;
                    background: #f8f9fa;
                }}
                .section {{
                    margin-bottom: 25px;
                }}
                .section-title {{
                    color: #333;
                    border-bottom: 2px solid #333;
                    padding-bottom: 5px;
                    margin-bottom: 15px;
                }}
                .experience-item {{
                    margin-bottom: 20px;
                }}
                .job-title {{
                    font-weight: bold;
                    font-size: 1.1em;
                }}
                .company {{
                    color: #666;
                    font-style: italic;
                }}
                .achievements {{
                    margin-top: 10px;
                }}
                .achievement {{
                    margin-left: 20px;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{personal.get('full_name', '')}</h1>
                <p>{personal.get('desired_position', '')}</p>
                <p>{personal.get('email', '')} | {personal.get('phone', '')} | {personal.get('location', '')}</p>
            </div>
        """
        
        # Professional Summary
        if cv_data.get("professional_summary"):
            html += f"""
            <div class="section">
                <h2 class="section-title">Professional Summary</h2>
                <p>{cv_data['professional_summary']}</p>
            </div>
            """
        
        # Work Experience
        if cv_data.get("work_experience"):
            html += """
            <div class="section">
                <h2 class="section-title">Work Experience</h2>
            """
            
            for exp in cv_data["work_experience"]:
                html += f"""
                <div class="experience-item">
                    <div class="job-title">{exp.get('job_title', '')}</div>
                    <div class="company">{exp.get('company', '')}</div>
                    <div class="dates">{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}</div>
                """
                
                if exp.get("achievements"):
                    html += '<div class="achievements">'
                    for achievement in exp["achievements"]:
                        html += f'<div class="achievement">• {achievement}</div>'
                    html += '</div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Education
        if cv_data.get("education"):
            html += """
            <div class="section">
                <h2 class="section-title">Education</h2>
            """
            
            for edu in cv_data["education"]:
                html += f"""
                <div class="education-item">
                    <div class="degree">{edu.get('degree', '')}</div>
                    <div class="institution">{edu.get('institution', '')}</div>
                    <div class="dates">{edu.get('start_date', '')} - {edu.get('end_date', '')}</div>
                </div>
                """
            
            html += '</div>'
        
        html += """
        </body>
        </html>
        """
        
        return html
    
    def get_supported_formats(self) -> List[Dict[str, Any]]:
        """Get list of supported export formats with capabilities"""
        
        formats = []
        
        for format_type in ExportFormat:
            format_info = {
                "format": format_type.value,
                "name": format_type.value.upper(),
                "available": True,
                "features": []
            }
            
            if format_type == ExportFormat.PDF:
                format_info["name"] = "PDF"
                format_info["available"] = HAS_WEASYPRINT
                format_info["features"] = ["print_ready", "professional", "styling", "pagination"]
                if not HAS_WEASYPRINT:
                    format_info["error"] = "WeasyPrint not installed"
            
            elif format_type == ExportFormat.DOCX:
                format_info["name"] = "Word Document"
                format_info["available"] = HAS_DOCX
                format_info["features"] = ["editable", "microsoft_office", "styling"]
                if not HAS_DOCX:
                    format_info["error"] = "python-docx not installed"
            
            elif format_type == ExportFormat.TXT:
                format_info["name"] = "Plain Text"
                format_info["features"] = ["portable", "lightweight", "ats_friendly"]
            
            elif format_type == ExportFormat.HTML:
                format_info["name"] = "HTML"
                format_info["features"] = ["web_ready", "styling", "portable"]
            
            elif format_type == ExportFormat.JSON:
                format_info["name"] = "JSON Data"
                format_info["features"] = ["data_portable", "structured", "machine_readable"]
            
            formats.append(format_info)
        
        return formats

# Global service instance
export_service = ExportService()